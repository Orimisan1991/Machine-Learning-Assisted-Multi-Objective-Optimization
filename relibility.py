# =========================================================
# COMPLETE SOCIAL MEDIA MARKETING ANALYSIS MODEL
# REGRESSION + DESCRIPTIVE + CORRELATION + RELIABILITY
# =========================================================

import pandas as pd
import numpy as np
import os
import re
import sys
import statsmodels.api as sm
from statsmodels.formula.api import ols
from scipy.stats import pearsonr
from docx import Document

# =========================================================
# OUTPUT FOLDER
# =========================================================

out = "Complete_Social_Media_Analysis_Output"
os.makedirs(out, exist_ok=True)

# =========================================================
# LOAD DATA
# =========================================================

df = pd.read_csv("QUESTION_6.csv")

# =========================================================
# CLEAN COLUMN NAMES
# =========================================================

df.columns = (
    df.columns.astype(str)
    .str.strip()
    .str.lower()
    .str.replace("\n", " ", regex=False)
    .str.replace('"', '', regex=False)
    .str.replace(r"\s+", " ", regex=True)
)

# =========================================================
# SAFE LOGGING
# =========================================================

log_file = open(f"{out}/analysis_output_log.txt", "w", encoding="utf-8")

class Logger:
    def write(self, message):
        try:
            sys.__stdout__.write(message)

            if log_file and not log_file.closed:
                log_file.write(message)

        except:
            pass

    def flush(self):
        pass

sys.stdout = Logger()

# =========================================================
# LIKERT SCALE CLEANER
# =========================================================

likert = {
    "strongly disagree": 1,
    "disagree": 2,
    "neutral": 3,
    "agree": 4,
    "strongly agree": 5
}

def clean_likert(x):

    if pd.isna(x):
        return np.nan

    x = str(x).strip().lower()

    x = re.sub(r"[^a-z ]", "", x)

    return likert.get(x, np.nan)

# =========================================================
# SMART COLUMN FINDER
# =========================================================

def find_column(df, keyword):

    keyword = keyword.lower().strip()

    for col in df.columns:

        if keyword in col:
            return col

    return None

# =========================================================
# QUESTION MAPPING
# =========================================================

Q_keywords = {

# =========================================================
# SOCIAL MEDIA PRESENCE
# =========================================================

"Q1": "active social media accounts",
"Q2": "regularly update our social media pages",
"Q3": "respond promptly to customer messages online",
"Q4": "social media platforms increase our visibility",
"Q5": "customers easily find our business through social media",

# =========================================================
# SOCIAL MEDIA ADVERTISING
# =========================================================

"Q6": "uses paid advertising on social media",
"Q7": "social media ads help us reach our target customers",
"Q8": "social media advertising increases our sales",
"Q9": "allocate part of our budget to online advertising",
"Q10": "social media advertising gives better results than traditional marketing",

# =========================================================
# CONTENT MARKETING
# =========================================================

"Q11": "post content regularly on our social media platforms",
"Q12": "content is engaging and attracts customers",
"Q13": "use images/videos to promote our products/services",
"Q14": "content reflects our brand identity clearly",
"Q15": "quality content helps us retain customers",

# =========================================================
# BUSINESS PERFORMANCE
# =========================================================

"Q16": "business performs better than competitors",
"Q17": "gained more customers through social media",
"Q18": "brand is well recognized in the market",
"Q19": "social media has improved our profitability",
"Q20": "maintain strong customer loyalty compared to competitors"
}

# =========================================================
# CREATE CLEAN VARIABLES
# =========================================================

print("\n================================================")
print("COLUMN MATCHING")
print("================================================")

for q, keyword in Q_keywords.items():

    col = find_column(df, keyword)

    if col is not None:

        df[q] = df[col].apply(clean_likert)

        print(f"{q}  --->  {col}")

    else:
        print(f"WARNING: {q} not found")

# =========================================================
# DEMOGRAPHIC ANALYSIS
# =========================================================

demo_cols = [
    "gender",
    "age",
    "position in business",
    "type of business",
    "years of operation"
]

print("\n================================================")
print("DEMOGRAPHIC ANALYSIS")
print("================================================")

for col in demo_cols:

    if col in df.columns:

        print(f"\n===== {col.upper()} =====")

        freq = df[col].value_counts(dropna=False)

        percent = (
            df[col]
            .value_counts(normalize=True, dropna=False)
            * 100
        )

        result = pd.DataFrame({
            "Frequency": freq,
            "Percentage": percent.round(2)
        })

        print(result)

        result.to_csv(f"{out}/{col}_demographics.csv")

# =========================================================
# QUESTION ANALYSIS
# =========================================================

print("\n================================================")
print("QUESTION ANALYSIS")
print("================================================")

for q in Q_keywords.keys():

    if q in df.columns:

        print(f"\n===== {q} =====")

        freq = df[q].value_counts(dropna=False).sort_index()

        percent = (
            df[q]
            .value_counts(normalize=True, dropna=False)
            .sort_index() * 100
        )

        mean_score = df[q].mean()

        std_score = df[q].std()

        result = pd.DataFrame({
            "Frequency": freq,
            "Percentage": percent.round(2)
        })

        print(result)

        print(f"Mean = {mean_score:.2f}")
        print(f"Std Dev = {std_score:.2f}")

        result.to_csv(f"{out}/{q}_summary.csv")

# =========================================================
# CREATE COMPOSITE VARIABLES
# =========================================================

presence_vars = ["Q1", "Q2", "Q3", "Q4", "Q5"]

advert_vars = ["Q6", "Q7", "Q8", "Q9", "Q10"]

content_vars = ["Q11", "Q12", "Q13", "Q14", "Q15"]

performance_vars = ["Q16", "Q17", "Q18", "Q19", "Q20"]

# Remove unavailable columns
presence_vars = [x for x in presence_vars if x in df.columns]
advert_vars = [x for x in advert_vars if x in df.columns]
content_vars = [x for x in content_vars if x in df.columns]
performance_vars = [x for x in performance_vars if x in df.columns]

# Compute means
df["social_media_presence"] = df[presence_vars].mean(axis=1)

df["social_media_advertising"] = df[advert_vars].mean(axis=1)

df["content_marketing"] = df[content_vars].mean(axis=1)

df["business_performance"] = df[performance_vars].mean(axis=1)

# =========================================================
# DESCRIPTIVE STATISTICS
# =========================================================

print("\n================================================")
print("DESCRIPTIVE STATISTICS")
print("================================================")

desc = df[[
    "social_media_presence",
    "social_media_advertising",
    "content_marketing",
    "business_performance"
]].describe()

print(desc)

desc.to_csv(f"{out}/descriptive_statistics.csv")

# =========================================================
# RELIABILITY TEST (CRONBACH ALPHA)
# =========================================================

def cronbach_alpha(items_df):

    items_df = items_df.dropna()

    item_scores = items_df.to_numpy()

    item_variance = item_scores.var(axis=0, ddof=1)

    total_score = item_scores.sum(axis=1)

    n_items = items_df.shape[1]

    total_variance = total_score.var(ddof=1)

    alpha = (
        n_items / (n_items - 1)
    ) * (
        1 - item_variance.sum() / total_variance
    )

    return alpha

print("\n================================================")
print("RELIABILITY TEST")
print("================================================")

alpha_presence = cronbach_alpha(df[presence_vars])

alpha_advert = cronbach_alpha(df[advert_vars])

alpha_content = cronbach_alpha(df[content_vars])

alpha_performance = cronbach_alpha(df[performance_vars])

print(f"Social Media Presence Alpha: {alpha_presence:.3f}")

print(f"Social Media Advertising Alpha: {alpha_advert:.3f}")

print(f"Content Marketing Alpha: {alpha_content:.3f}")

print(f"Business Performance Alpha: {alpha_performance:.3f}")

# =========================================================
# CORRELATION ANALYSIS
# =========================================================

print("\n================================================")
print("CORRELATION ANALYSIS")
print("================================================")

corr_data = df[[
    "social_media_presence",
    "social_media_advertising",
    "content_marketing",
    "business_performance"
]]

corr_matrix = corr_data.corr()

print(corr_matrix)

corr_matrix.to_csv(f"{out}/correlation_matrix.csv")

# =========================================================
# MULTIPLE REGRESSION MODEL
# =========================================================

print("\n================================================")
print("MULTIPLE REGRESSION ANALYSIS")
print("================================================")

IVs = [
    "social_media_presence",
    "social_media_advertising",
    "content_marketing"
]

DV = "business_performance"

reg_df = df[IVs + [DV]].dropna()

if len(reg_df) > 5:

    # =====================================================
    # FIT MODEL
    # =====================================================

    X = sm.add_constant(reg_df[IVs])

    y = reg_df[DV]

    model = sm.OLS(y, X).fit()

    # =====================================================
    # REGRESSION SUMMARY
    # =====================================================

    print(model.summary())

    # =====================================================
    # ANOVA
    # =====================================================

    formula = DV + " ~ " + " + ".join(IVs)

    anova_model = ols(formula, data=reg_df).fit()

    anova_table = sm.stats.anova_lm(anova_model, typ=2)

    print("\n================================================")
    print("ANOVA TABLE")
    print("================================================")

    print(anova_table)

    # =====================================================
    # COEFFICIENT TABLE
    # =====================================================

    coef = model.summary2().tables[1]

    print("\n================================================")
    print("COEFFICIENT TABLE")
    print("================================================")

    print(coef)

    # =====================================================
    # SAVE REGRESSION OUTPUTS
    # =====================================================

    coef.to_csv(f"{out}/regression_coefficients.csv")

    anova_table.to_csv(f"{out}/anova_table.csv")

# =========================================================
# HYPOTHESIS TESTING
# =========================================================

print("\n================================================")
print("HYPOTHESIS TESTING")
print("================================================")

for variable in IVs:

    p_value = model.pvalues[variable]

    if p_value < 0.05:

        print(f"{variable}: SIGNIFICANT (p = {p_value:.4f})")

    else:

        print(f"{variable}: NOT SIGNIFICANT (p = {p_value:.4f})")

# =========================================================
# SAVE CLEANED DATA
# =========================================================

df.to_csv(f"{out}/cleaned_dataset.csv", index=False)

# =========================================================
# WORD REPORT
# =========================================================

doc = Document()

doc.add_heading(
    "Social Media Marketing and Business Performance Analysis",
    1
)

# =========================================================
# DESCRIPTIVE STATISTICS
# =========================================================

doc.add_heading("Descriptive Statistics", level=2)

doc.add_paragraph(desc.to_string())

# =========================================================
# RELIABILITY
# =========================================================

doc.add_heading("Reliability Test", level=2)

doc.add_paragraph(
    f"Social Media Presence Alpha = {alpha_presence:.3f}"
)

doc.add_paragraph(
    f"Social Media Advertising Alpha = {alpha_advert:.3f}"
)

doc.add_paragraph(
    f"Content Marketing Alpha = {alpha_content:.3f}"
)

doc.add_paragraph(
    f"Business Performance Alpha = {alpha_performance:.3f}"
)

# =========================================================
# CORRELATION
# =========================================================

doc.add_heading("Correlation Matrix", level=2)

doc.add_paragraph(corr_matrix.to_string())

# =========================================================
# REGRESSION
# =========================================================

doc.add_heading("Regression Result", level=2)

doc.add_paragraph(model.summary().as_text())

# =========================================================
# ANOVA
# =========================================================

doc.add_heading("ANOVA Table", level=2)

doc.add_paragraph(anova_table.to_string())

# =========================================================
# COEFFICIENTS
# =========================================================

doc.add_heading("Coefficient Table", level=2)

doc.add_paragraph(coef.to_string())

# =========================================================
# SAVE REPORT
# =========================================================

doc.save(f"{out}/Complete_Analysis_Report.docx")

# =========================================================
# FINAL OUTPUT
# =========================================================

print("\n================================================")
print("ANALYSIS COMPLETED SUCCESSFULLY")
print("================================================")

print("FILES SAVED IN:", out)

log_file.close()